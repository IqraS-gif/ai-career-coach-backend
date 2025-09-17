from datetime import datetime # Correct import: datetime now refers to the datetime.datetime class
import os
import io
import sys
import json
import re
from typing import Optional, Tuple, List, Dict, Any, Union

# Required libraries (ensure they are installed via requirements.txt)
import fitz  # PyMuPDF
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =========================
# Setup
# =========================
def setup_api():
    """Loads environment variables and configures the API key."""
    load_dotenv()
    try:
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key: raise ValueError("GOOGLE_API_KEY not found in your .env file.")
        genai.configure(api_key=api_key)
        return "gemini-1.5-flash-latest"
    except Exception as e:
        print(f"Error: API configuration failed. {e}")
        sys.exit(1)

MODEL_NAME = setup_api()

# =========================
# JSON Schema Constants
# =========================

ASSESSMENT_QUESTIONS_SCHEMA = """
[
  {
    "question_id": "string",
    "question_text": "string",
    "question_type": "single_choice" | "multiple_choice" | "short_answer" | "coding_challenge",
    "options": ["string option 1", "string option 2", "string option 3", "string option 4"],
    "correct_answer_keys": ["string option 1"]
  }
]
"""

ASSESSMENT_EVALUATION_SCHEMA = """
{
  "overall_score": 75,
  "skills_mastered": 3,
  "areas_to_improve": 2,
  "skill_scores": { "Python": 80, "SQL": 60, "Data Analysis": 75 },
  "strengths": ["Demonstrated strong foundational knowledge in Python.", "Understood basic SQL queries."],
  "weaknesses": ["Struggled with complex data manipulation in SQL.", "Limited understanding of advanced data analysis concepts."],
  "recommendations": [
    "Focus on SQL subqueries and window functions for data manipulation.",
    "Practice implementing machine learning algorithms from scratch.",
    "Explore advanced data visualization techniques and tools."
  ]
}
"""

FULL_RESUME_ANALYSIS_SCHEMA = """
{
  "analysis_date": "September 05, 2025",
  "job_role_context": "string",  # CORRECTED: Changed default to 'string' to indicate it's dynamic
  "ai_model": "Google Gemini",
  "overall_resume_score": 68,
  "overall_resume_grade": "Good",
  "ats_optimization_score": 60,
  "professional_profile_analysis": {
    "title": "Professional Profile Analysis",
    "summary": "The candidate presents a clear trajectory of learning and project involvement, demonstrating foundational skills relevant to the target role. However, the profile could benefit from a more concise and impact-driven summary statement tailored directly to a 'Frontend Developer' role, immediately highlighting key value propositions."
  },
  "education_analysis": {
    "title": "Education Analysis",
    "summary": "The education section is concise but could be enhanced. Adding the expected graduation date would be helpful. Listing relevant development coursework (e.g., 'Web Development I & II', 'Data Structures and Algorithms') would reinforce technical expertise. Emphasize any honors or significant academic achievements."
  },
  "experience_analysis": {
    "title": "Experience Analysis",
    "summary": "The projects section is informative but overwhelming. The descriptions are too long and lack quantifiable achievements. Instead of lengthy descriptions, focus on impactful results using numbers and action verbs. For instance, 'Developed a responsive e-commerce platform that increased user engagement by 15%.'"
  },
  "skills_analysis": {
    "title": "Skills Analysis",
    "summary": "The current skills section is adequate but could be structured more effectively for ATS scanning. Consider grouping related skills (e.g., 'Languages: JavaScript, Python', 'Frameworks: React, Angular'). Ensure all frontend-specific skills (e.g., HTML5, CSS3, SASS, Webpack) are explicitly listed and visible."
  },
  "key_strengths": [
    "Diverse Project Portfolio: The candidate has undertaken a wide variety of projects, showcasing initiative and a broad skillset.",
    "Clear Learning Trajectory: Demonstrates continuous learning and application of new technologies.",
    "Foundational Technical Acumen: Possesses a solid understanding of core computer science principles applicable to development."
  ],
  "areas_for_improvement": [
    "Target Role Focus: The resume needs to be sharply focused on the frontend developer role. De-emphasize or remove projects that don't directly showcase relevant frontend skills.",
    "Quantify Achievements: Introduce more quantifiable results (numbers, percentages) in project and experience descriptions.",
    "ATS Keyword Optimization: Integrate common frontend developer keywords (e.g., 'Responsive Design', 'API Integration', 'Version Control', 'TypeScript') more strategically.",
    "Concise Descriptions: Shorten lengthy project/experience descriptions to impactful bullet points.",
    "Consistent Formatting: Ensure consistent formatting across all sections, especially dates and bullet points, for better readability and ATS parsing."
  ],
  "overall_assessment": "This resume demonstrates a strong foundation in computer engineering and relevant project involvement. With targeted refinement to highlight frontend development skills, quantify achievements, and optimize for ATS, the candidate can significantly enhance their chances of securing interviews."
}
"""


# =========================
# Helper Functions
# =========================
def _safe_json_loads(s: str, fallback=None):
    """Safely loads a JSON string, even if it's embedded in markdown."""
    if not s: return fallback
    s = s.strip()
    # Attempt to strip markdown code blocks
    if s.startswith("```json"): s = s[7:]
    if s.endswith("```"): s = s[:-3]
    s = s.strip() # Strip again after removing markdown

    try:
        return json.loads(s)
    except json.JSONDecodeError:
        # If direct parse fails, try to find a JSON-like substring
        m = re.search(r"\{.*\}", s, flags=re.DOTALL)
        if m:
            try: return json.loads(m.group(0))
            except json.JSONDecodeError: return fallback
    return fallback

def _norm(s: Optional[str]) -> bool:
    """Returns True if the stripped string is not empty, False otherwise."""
    return bool(s and s.strip())

def _smart_join(parts: List[Optional[str]]) -> str:
    """Joins a list of parts with a separator, ignoring empty or None parts."""
    return " | ".join([str(p) for p in parts if _norm(p)])

def _best_section_key(target_key: str, available_keys: List[str]) -> Optional[str]:
    """Finds the best matching key in a dictionary from a fuzzy user input."""
    if not target_key: return None
    t = target_key.strip().lower().replace(" ", "_").replace("-", "_")
    for k in available_keys:
        k_norm = k.lower().replace(" ", "_")
        if t == k_norm or t in k_norm or k_norm in t: return k
    return None

def parse_user_optimization_input(inp: str) -> Tuple[Optional[str], Optional[str]]:
    """Parses user input into a (section, instruction) tuple."""
    val = (inp or "").strip()
    if not val: return None, None
    if ":" in val:
        left, right = val.split(":", 1); return _norm(left), _norm(right)
    if len(val.split()) == 1:
        return val, None
    return None, val

def _stringify_list_content(content: Any) -> str:
    """Safely converts a list of strings or dicts into a single newline-separated string."""
    if not isinstance(content, list): return str(content or "")
    string_parts = []
    for item in content:
        if isinstance(item, str): string_parts.append(item)
        elif isinstance(item, dict):
            string_parts.append(", ".join([f"{k.replace('_', ' ').title()}: {v}" for k, v in item.items()]))
        else: string_parts.append(str(item))
    return "\n".join(string_parts)

def extract_text_auto(file_content: bytes, file_extension: str) -> Optional[str]:
    """
    Automatically detects file type and extracts text from file content (bytes).
    Returns extracted text or None if extraction fails.
    """
    print(f"DEBUG(ai_core): extract_text_auto called for in-memory content (Type: {file_extension})")
    
    try:
        if file_extension == ".pdf":
            print(f"DEBUG(ai_core): Attempting PDF extraction using fitz from bytes.")
            with fitz.open(stream=file_content, filetype="pdf") as doc: 
                text_content = "\n".join([page.get_text() for page in doc])
                print(f"DEBUG(ai_core): PDF extraction finished. Length: {len(text_content)}")
                return text_content
        elif file_extension == ".docx":
            print(f"DEBUG(ai_core): Attempting DOCX extraction using python-docx from bytes.")
            doc = Document(io.BytesIO(file_content))
            
            all_paragraphs = []
            for i, p in enumerate(doc.paragraphs):
                p_text = p.text
                all_paragraphs.append(p_text)
            
            print(f"DEBUG(ai_core): Found {len(all_paragraphs)} raw paragraphs in DOCX.")
            
            sample_paragraphs = [p for p in all_paragraphs if _norm(p)][:5]
            if sample_paragraphs:
                print(f"DEBUG(ai_core): Sample non-empty paragraphs (first {len(sample_paragraphs)}):")
                for i, p_text in enumerate(sample_paragraphs):
                    print(f"  [{i+1}] '{p_text[:70]}...'")
            else:
                print(f"DEBUG(ai_core): No non-empty paragraphs found initially by python-docx.")

            chunks = [p_text for p_text in all_paragraphs if _norm(p_text)]
            print(f"DEBUG(ai_core): {len(chunks)} paragraphs passed _norm filter.")

            if doc.tables:
                print(f"DEBUG(ai_core): Found {len(doc.tables)} tables in DOCX.")
                for table_idx, table in enumerate(doc.tables):
                    print(f"DEBUG(ai_core): Processing table {table_idx + 1}.")
                    for row_idx, row in enumerate(table.rows):
                        row_cells_text = []
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = cell.text
                            row_cells_text.append(cell_text)
                        
                        cells_for_chunk = [text for text in row_cells_text if _norm(text)]
                        if cells_for_chunk: 
                            joined_cells = " | ".join(cells_for_chunk)
                            chunks.append(joined_cells)
                            print(f"DEBUG(ai_core): Added row {row_idx + 1} from table {table_idx + 1}. Content: '{joined_cells[:70]}...'")
                else:
                    print(f"DEBUG(ai_core): No tables found in DOCX.")
            
            final_text = "\n".join(chunks)
            print(f"DEBUG(ai_core): DOCX extraction finished. Final text length: {len(final_text)}")
            return final_text
        else:
            print(f"ERROR(ai_core): Unsupported file type '{file_extension}'."); return None
    except Exception as e:
        print(f"ERROR(ai_core): Failed to read file content (type: {file_extension}). Exception: {e}", exc_info=True)
        return None

def get_resume_structure(resume_text: str) -> Optional[Dict[str, Any]]:
    prompt = f"""
You are an expert HR Technology engineer specializing in resume data extraction. Your task is to convert the raw text of a resume into a structured, valid JSON object, capturing ALL information with high fidelity.
**Instructions:**
1.  **Use the Base Schema:** For common sections, use the following schema.
2.  **Capture Everything Else:** If you find other sections that do not fit the schema (e.g., "Achievements", "Leadership"), create a new top-level key for them (e.g., "achievements").
3.  **IGNORE THE SKILLS SECTION:** Do not parse the skills section in this step. It will be handled by a different process. Omit the 'skills' key from your output.
**Base Schema:**
{{
  "personal_info": {{ "name": "string", "email": "string", "phone": "string", "linkedin": "string", "github": "string" }},
  "summary": "string",
  "work_experience": [ {{ "role": "string", "company": "string", "duration": "string", "description": ["string", ...] }} ],
  "internships": [ {{ "role": "string", "company": "string", "duration": "string", "description": ["string", ...] }} ],
  "education": [ {{ "institution": "string", "degree": "string", "duration": "string", "description": ["string", ...] }} ],
  "projects": [ {{ "title": "string", "description": ["string", ...] }} ],
  "certifications": [ {{ "name": "string", "description": "string" }} ]
}}
**Critical Rules:**
- If a section from the base schema is NOT in the resume, YOU MUST OMIT ITS KEY from the final JSON. Do not create empty sections.
- Your final output must be a single, valid JSON object starting with `{{` and ending with `}}`. Do not include markdown.
--- RESUME TEXT ---
{resume_text}
--- END RESUME TEXT ---
"""
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        data = _safe_json_loads(response.text, fallback=None)
        if not data:
            print("\n--- ERROR: GEMINI API FAILED TO RETURN VALID JSON (STRUCTURE) ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-----------------------------------------------------\n")
            return None
        return data
    except Exception as e:
        print(f"Error during API call to Gemini: {e}"); return None

def categorize_skills_from_text(resume_text: str) -> Optional[Dict[str, List[str]]]:
    prompt = f"""
You are an expert technical recruiter and data analyst.
Your sole job is to scan the entire resume text provided and identify all skills, both technical and soft.
**Instructions:**
1.  Extract skills from *anywhere* in the text: summaries, project descriptions, a dedicated skills section, etc.
2.  Categorize the skills into the predefined keys in the JSON schema below.
3.  Place each skill only in the most appropriate category.
4.  If a category has no skills, you can omit the key from the output.
**JSON Output Schema:**
{{
    "Programming Languages": ["Python", "JavaScript", "Java", "C++", ...],
    "Frameworks and Libraries": ["TensorFlow", "PyTorch", "React", "Node.js", "Pandas", ...],
    "Databases": ["MySQL", "PostgreSQL", "MongoDB", ...],
    "Tools and Platforms": ["Git", "Docker", "AWS", "Jira", "Linux", ...],
    "Data Science": ["Machine Learning", "NLP", "Data Visualization", "Predictive Modeling", ...],
    "Soft Skills": ["Leadership", "Teamwork", "Communication", "Problem Solving", ...]
}}
**Critical Rules:**
- Your output must be ONLY the valid JSON object described above.
- Do not add any explanation or markdown.
--- RESUME TEXT ---
{resume_text}
--- END RESUME TEXT ---
"""
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        data = _safe_json_loads(response.text, fallback=None)
        if not data:
            print("\n--- ERROR: GEMINI FAILED TO INFER SKILLS ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-------------------------------------------------\n")
            return None
        return data
    except Exception as e:
        print(f"Error inferring skills with Gemini: {e}"); return None

def optimize_resume_json(resume_json: Dict[str, Any], user_input: str, job_description: Optional[str] = None) -> Dict[str, Any]:
    section_req, instruction = parse_user_optimization_input(user_input)
    keys_present = list(resume_json.keys())
    model = genai.GenerativeModel(MODEL_NAME)

    # --- NEW DEBUGGING BLOCK (CRITICAL FOR THIS ERROR) ---
    print(f"DEBUG(ai_core - optimize_resume_json): Verifying resume_json for non-serializable types before AI prompt.")
    def check_for_non_serializable(obj: Any, path: str = "") -> bool:
        if isinstance(obj, dict):
            for k, v in obj.items():
                new_path = f"{path}.{k}" if path else k
                if check_for_non_serializable(v, new_path):
                    return True
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                new_path = f"{path}[{i}]"
                if check_for_non_serializable(item, new_path):
                    return True
        elif isinstance(obj, datetime): # CORRECTED: Use datetime.datetime for explicit type check
            print(f"ERROR(ai_core - optimize_resume_json): Found UNCONVERTED datetime object at path: '{path}'. Value: {obj}")
            return True
        return False
    
    if check_for_non_serializable(resume_json):
        print("CRITICAL ERROR: Unconverted datetime type detected in resume_json. The `_convert_firestore_timestamps` function in `db_core.py` is not fully converting data from Firestore or being applied correctly.")
        raise TypeError("Resume data contains non-JSON serializable datetime objects. Check db_core.py's _convert_firestore_timestamps.")
    # --- END NEW DEBUGGING BLOCK ---

    job_desc_context = ""
    if job_description and job_description.strip():
        job_desc_context = f"""
        **Job Description Context:**
        Below is the job description for which the resume is being optimized. Incorporate keywords, desired skills, and align the achievements to the requirements of this role.
        ```
        {job_description}
        ```
        """

    base_prompt_context = f"""
CONTEXT: You are an elite career strategist and executive resume writer. Your task is to transform a resume from a passive list of duties into a compelling narrative of achievements that will impress top-tier recruiters.
**Your Transformation Checklist (Apply to every relevant bullet point):**
1.  **Lead with a Powerful Action Verb:** Replace weak verbs with strong, specific verbs (e.g., "Engineered," "Architected," "Spearheaded").
2.  **Quantify Metrics Relentlessly:** Add concrete numbers to show scale and achievement.
3.  **Showcase Impact and Scope:** If a number isn't available, describe the tangible impact or business outcome.
4.  **Integrate Technical Skills Naturally:** Weave technologies into the story of the achievement.
5.  **Ensure Brevity and Clarity:** Remove filler words. Each bullet point should be a single, powerful line.

{job_desc_context} 

**Critical Rules:**
- **Do not modify, add, or delete any titles, names, companies, institutions, or skill names.** This is a strict rule. Only rewrite descriptions.
- DO NOT invent facts or skills.
- DO NOT invent specific numbers.
- Preserve the original data structure.
- Do not modify personal information (name, email, phone).
- Your final output must be only the requested, valid JSON. Do not include markdown.
"""
    if section_req:
        mapped = _best_section_key(section_req, keys_present)
        if not mapped:
            print(f"⚠️ Section '{section_req}' not found."); return resume_json
        sec_data = resume_json.get(mapped)
        instr_text = instruction or "Apply your transformation checklist to make this section world-class."
        prompt = f"""
{base_prompt_context}
TASK: Apply your full transformation checklist to optimize ONLY the following JSON section, named "{mapped}".
--- INPUT JSON SECTION ---
{json.dumps(sec_data, indent=2)}
--- END INPUT JSON ---
"""
    else:
        prompt = f"""
{base_prompt_context}
TASK: Apply your full transformation checklist to optimize all sections of the following resume JSON.
--- FULL INPUT JSON ---
{json.dumps(resume_json, indent=2)}
--- END INPUT JSON ---
"""
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        optimized_data = _safe_json_loads(response.text, fallback=None)
        
        if not optimized_data:
            print("\n--- ERROR: GEMINI API FAILED TO RETURN VALID JSON (OPTIMIZE) ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-----------------------------------------------------\n")
            return resume_json
            
        if section_req and optimized_data:
            resume_json[mapped] = optimized_data
        elif optimized_data:
            for key, value in optimized_data.items():
                if key in resume_json and isinstance(resume_json[key], dict) and isinstance(value, dict):
                    resume_json[key].update(value)
                elif key in resume_json and isinstance(resume_json[key], list) and isinstance(value, list):
                    resume_json[key] = value
                else:
                    resume_json[key] = value
    except Exception as e:
        print(f"Error during optimization: {e}")
    return resume_json

def optimize_for_linkedin(resume_json: Dict[str, Any], user_input: str, job_description: Optional[str] = None) -> Optional[Dict[str, Any]]:
    context_text = []
    if 'summary' in resume_json: context_text.append(f"Summary:\n{resume_json['summary']}")
    
    all_experiences = resume_json.get('work_experience', []) + resume_json.get('internships', [])
    if all_experiences:
        context_text.append("\nProfessional Experience & Internships:")
        for job in all_experiences:
            description_str = ' '.join(job.get('description', []) if isinstance(job.get('description'), list) else [str(job.get('description', ''))])
            context_text.append(f"- {job.get('role')} at {job.get('company')}: {description_str}")
    
    if 'projects' in resume_json:
        context_text.append("\nProjects:")
        for project in resume_json['projects']:
            description_str = ' '.join(project.get('description', []) if isinstance(project.get('description'), list) else [str(project.get('description', ''))])
            context_text.append(f"- {project.get('title')}: {description_str}")
    
    if 'skills' in resume_json and isinstance(resume_json['skills'], dict):
        skills_summary = ", ".join([f"{cat}: {', '.join(skills)}" for cat, skills in resume_json['skills'].items()])
        context_text.append(f"\nSkills: {skills_summary}")
    
    for key, value in resume_json.items():
        if key not in ['personal_info', 'summary', 'work_experience', 'internships', 'projects', 'skills', 'education', 'certifications', 'resume_metadata', 'raw_text']:
            if isinstance(value, str):
                context_text.append(f"\n{key.replace('_', ' ').title()}:\n{value}")
            elif isinstance(value, list):
                context_text.append(f"\n{key.replace('_', ' ').title()}:\n" + "\n".join([str(item) for item in value]))

    resume_context = "\n".join(context_text)
    section_req, instruction = parse_user_optimization_input(user_input)

    job_desc_context = ""
    if job_description and job_description.strip():
        job_desc_context = f"""
        **Job Description Context:**
        Below is the job description for which the LinkedIn profile is being optimized. Align the content with the keywords, requirements, and tone of this role.
        ```
        {job_description}
        ```
        """

    base_prompt_context = f"""
You are an expert LinkedIn profile strategist and personal branding coach.
Your task is to generate compelling, optimized text for a user's LinkedIn profile based on the provided resume content.
**Instructions:**
1.  **Headlines:** Create 2-3 powerful, keyword-rich headline options.
2.  **About (Summary):** Write a compelling, first-person "About" section.
3.  **Experiences:** For EACH job/internship in the context, rewrite the bullet points to be concise and results-oriented.
4.  **Projects:** For EACH project in the context, rewrite its description to be engaging for a LinkedIn audience.

{job_desc_context}
**JSON Output Schema:**
{{
    "headlines": ["string option 1", ...],
    "about_section": "string",
    "optimized_experiences": [ {{ "title": "Role at Company", "description": "string" }} ],
    "optimized_projects": [ {{ "title": "Project Title", "description": "string" }} ]
}}

**Critical Rules:**
- Generate content ONLY from the provided resume context.
- Keep the tone professional but approachable.
- Your final output must be ONLY the valid JSON object that matches the requested task.
"""
    if section_req:
        instr_text = instruction or f"Make the {section_req} section more compelling and professional."
        prompt = f"""
{base_prompt_context}
TASK: Based on the resume context, optimize ONLY the '{section_req}' portion of a LinkedIn profile.
--- RESUME CONTEXT ---
{resume_context}
--- END RESUME CONTEXT ---
"""
    else:
        instr_text = instruction or "Optimize the entire LinkedIn profile, processing every experience and project."
        prompt = f"""
{base_prompt_context}
TASK: Based on the resume context, perform a full optimization of a LinkedIn profile.
--- RESUME CONTEXT ---
{resume_context}
--- END RESUME CONTEXT ---
"""
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        data = _safe_json_loads(response.text, fallback=None)
        if not data:
            print("\n--- ERROR: GEMINI FAILED TO INFER LINKEDIN CONTENT ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-------------------------------------------------\n")
            return None
        return data
    except Exception as e:
        print(f"Error generating LinkedIn content with Gemini: {e}"); return None

def save_resume_json_to_docx(resume_json: Dict[str, Any]) -> Document:
    doc = Document()
    def add_heading(text: Optional[str], level: int = 1):
        t = (text or "").strip(); 
        if t: h = doc.add_heading(t, level=level); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def add_para(text: Optional[str], bold: bool = False, style: Optional[str] = None):
        t = (text or "").strip() 
        if t:
            p = doc.add_paragraph(style=style)
            run = p.add_run(t)
            run.bold = bold
            run.font.size = Pt(11)
            if style == "List Bullet": p.paragraph_format.left_indent = Pt(36)
            
    print_order = ['personal_info', 'summary', 'skills', 'work_experience', 'internships', 'projects', 'education', 'certifications']
    
    name_for_title = resume_json.get('personal_info', {}).get('name', '')
    if name_for_title:
        doc.add_heading(name_for_title, level=0)

    contact_info_parts = []
    p_info = resume_json.get('personal_info', {})
    if p_info.get('email'): contact_info_parts.append(p_info['email'])
    if p_info.get('phone'): contact_info_parts.append(p_info['phone'])
    if p_info.get('linkedin'): contact_info_parts.append(p_info['linkedin'])
    if p_info.get('github'): contact_info_parts.append(p_info['github'])
    if contact_info_parts:
        add_para(_smart_join(contact_info_parts))
    
    for section in print_order:
        if section in resume_json:
            content = resume_json[section]
            if section == 'personal_info':
                continue
            
            add_heading(section.replace("_", " ").title(), level=2)
            
            if section == 'summary' and isinstance(content, str):
                add_para(content)
            elif section == 'skills' and isinstance(content, dict):
                for category, skill_list in content.items():
                    if isinstance(skill_list, list) and skill_list:
                        p = doc.add_paragraph();
                        run = p.add_run(category.replace("_", " ").title() + ': '); run.bold = True
                        p.add_run(", ".join(skill_list)); 
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, str):
                        add_para(item, style="List Bullet")
                    elif isinstance(item, dict):
                        # Ensure any dates/timestamps within item are converted to string before joining/displaying
                        item_copy = item.copy()
                        for k, v in item_copy.items():
                            if isinstance(v, datetime): # CORRECTED: Check for datetime.datetime objects
                                item_copy[k] = v.strftime("%b %d, %Y") # Format as readable string
                        
                        header_parts = [item_copy.get("title"), item_copy.get("name"), item_copy.get("role"), item_copy.get("degree"), item_copy.get("institution")]
                        header = _smart_join(header_parts)
                        if header: add_para(header, bold=True)
                        
                        sub_header_parts = [item_copy.get("company"), item_copy.get("duration")]
                        sub_header = _smart_join(sub_header_parts)
                        if sub_header: add_para(sub_header)
                        
                        desc = item_copy.get("description", [])
                        if isinstance(desc, list):
                            for bullet in desc:
                                if _norm(bullet): add_para(str(bullet), style="List Bullet")
                        elif isinstance(desc, str) and _norm(desc):
                            add_para(str(desc), style="List Bullet")

            elif isinstance(content, str):
                add_para(content)
            
    for section, content in resume_json.items():
        if section not in print_order and section not in ['resume_metadata', 'raw_text','optimized_summary']:
            add_heading(section.replace("_", " ").title(), level=2)
            if isinstance(content, list):
                for item in content: add_para(str(item), style="List Bullet")
            else: add_para(str(content))
            
    print("\n✅ DOCX document generated in memory.")
    return doc

def generate_career_roadmap(user_profile: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    prompt = f"""
    Act as a world-class AI Career Strategist and Technical Project Manager. Your task is to generate a deeply personalized, multi-faceted career action plan.

    **STEP 1: ANALYZE THE USER'S PROFILE**
    - **User's Current State (from resume or manual input):** ```{user_profile.get('current_skills_input')}```
    - **Stated Current Proficiency:** {user_profile.get('current_level')}
    - **User's Stated Goal (Job Description or Desired Skills):** ```{user_profile.get('goal_input')}```
    - **Desired Goal Proficiency:** {user_profile.get('goal_level')}
    - **Time Commitment:** Plan for a duration of **{user_profile.get('duration')}**, assuming **{user_profile.get('study_hours')}** study hours per month.

    **STEP 2: GENERATE THE ACTION PLAN AS A SINGLE, VALID JSON OBJECT**
    The JSON output must be perfectly structured with the following keys. Do not include any explanatory text outside of the JSON object.

    1.  "domain": A single string representing the most relevant domain inferred from the goal input (e.g., "Data Science", "Cybersecurity"). This is a new, crucial key.
    2.  "extracted_skills_and_projects": A JSON object with "skills" (array of strings) and "projects" (array of strings).
    3.  "job_match_score": A JSON object with "score" (number) and "summary" (string).
    4.  "skills_to_learn_summary": An array of strings.
    5.  "timeline_chart_data": A JSON object with "labels" (array of strings) and "durations" (array of numbers in weeks) ans also the total weeks counts must be equal to users specified duration.
    6.  "detailed_roadmap": An array of "phase" objects, each with "phase_title", "phase_duration", and "topics" (array of strings).
    7.  "suggested_projects": An array of 2 "project" objects, each with "project_title", "project_level", "skills_mapped", "what_you_will_learn", and a multi-step "implementation_plan".
    8.  "suggested_courses": THIS IS A CRITICAL SECTION.
        - You MUST generate an array of 2-3 "course" objects.
        - Each object MUST contain the following FOUR keys: "course_name", "platform", "url", and "mapping".
        - The "platform" MUST be a string like "Coursera", "edX", "Pluralsight", etc.
        - The "url" MUST be a direct, fully-qualified, and workable hyperlink.
        - The "mapping" MUST be a concise sentence explaining how the course helps the roadmap.
        - **Follow this example format precisely:**
          `{{ "course_name": "Google Data Analytics Certificate", "platform": "Coursera", "url": "https://www.coursera.org/professional-certificates/google-data-analytics", "mapping": "This certificate covers the foundational skills in Phase 1 and 2." }}`
    """
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        cleaned_response_text = response.text.replace('```json', '').replace('```', '').strip()
        return json.loads(cleaned_response_text)
    except Exception as e:
        print(f"An error occurred during AI roadmap generation: {e}"); return None

def get_tutor_explanation(topic: str) -> Optional[Dict[str, Any]]:
    prompt = f"""
    Act as a friendly and encouraging expert tutor. A user is currently working through a personalized learning plan and is stuck on the following topic: **"{topic}"**

    Your task is to provide a clear, helpful explanation in a structured JSON format. The JSON object must have the following keys:

    1.  **"analogy"**: A simple, real-world analogy to help the user understand the core concept intuitively.
    2.  **"technical_definition"**: A concise, technically accurate definition. If the topic involves code, provide a short, well-commented code snippet in the appropriate language (e.g., Python, JavaScript).
    3.  **"prerequisites"**: An array of 1-3 prerequisite concepts the user might need to review. This helps them identify foundational knowledge gaps.

    Generate the JSON object and nothing else.
    """
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        cleaned_response_text = response.text.replace('```json', '').replace('```', '').strip()
        return json.loads(cleaned_response_text)
    except Exception as e:
        print(f"An error occurred in AI Tutor: {e}"); return None

# --- CORRECTED CHATBOT FUNCTION ---
def get_chatbot_response(query: str, history: list, career_plan_summary: str) -> dict:
    """
    Generates a chatbot response using the pre-summarized career plan string as context.
    
    Args:
        query: The user's latest question.
        history: The previous conversation history.
        career_plan_summary: A PRE-SUMMARIZED STRING of the user's career plan.
    """
    print("AI Core: Received request. The career plan context is a string.")
    
    system_prompt = (
        f"You are an AI career strategist and tutor. Your purpose is to provide concise, point-to-point, and beginner-friendly guidance to the user, strictly based on the career plan provided below.\n\n"
        f"**Career Plan Details:**\n{career_plan_summary}\n\n"
        f"**Your Instructions:**\n"
        f"1. Keep responses brief, beginner-friendly, and to the point.\n"
        f"2. You can answer questions related to the provided career plan, including the **job match score, priority skills, timeline, detailed roadmap, projects, and courses**.\n"
        f"3. If the user asks a question that is **outside the scope** of the career plan's domain or is not directly related to the provided plan data, you must respond with a polite refusal. For example, 'That question seems to be outside the scope of your current career plan. Is there anything I can help you with related to your career plan?'\n\n"
        f"Let's begin."
    )

    model_history = []
    for message in history:
        # Ensure role is either 'user' or 'model' as required by the API
        role = 'user' if message.get('role') == 'user' else 'model'
        content = message.get('content', '')
        # Prevent sending empty content parts
        if content:
            model_history.append({'role': role, 'parts': [content]})

    try:
        model = genai.GenerativeModel(MODEL_NAME)
        chat_session = model.start_chat(history=model_history)
        
        # We send the system prompt along with the query to ensure context is always present
        full_prompt = f"{system_prompt}\n\nUSER QUESTION: {query}"
        
        response = chat_session.send_message(full_prompt)
        
        if not response.text:
             return {"response": "I'm sorry, I couldn't generate a response for that. Please try rephrasing your question."}

        return {"response": response.text}

    except Exception as e:
        print(f"❌ ERROR in AI Core (get_chatbot_response): {e}")
        # Propagate the error up to be handled by the router
        raise e

def generate_assessment_questions(
    assessment_type: str,
    skills: List[str],
    target_role: Optional[str] = None,
    num_questions: int = 5,
    user_id: Optional[str] = None
) -> Optional[Dict[str, Any]]:
    """
    Generates a set of assessment questions based on selected skills and target role.
    Uses Gemini Flash.
    """
    
    skills_str = ", ".join(skills)
    role_context = f" for a {target_role}" if target_role else ""

    difficulty_hint = "medium difficulty"
    normalized_role = (target_role or "").lower()
    if "junior" in normalized_role:
        difficulty_hint = "beginner to medium difficulty"
    elif "senior" in normalized_role or "lead" in normalized_role:
        difficulty_hint = "medium to advanced difficulty"


    prompt = f"""
    You are an expert technical interviewer and AI assessment designer.
    Your task is to generate a concise, focused skill assessment with exactly {num_questions} questions.
    The assessment should cover the following skills: **{skills_str}**.
    The target context is {assessment_type.replace('_', ' ').title()} role{role_context}, at a {difficulty_hint} level.

    **Instructions for Question Generation:**
    1.  Generate a mix of question types:
        -   **Single-choice (radio buttons):** ~50% of questions. Provide 4 distinct options.
        -   **Multiple-choice (checkboxes):** ~20% of questions. Provide 4 distinct options, clearly indicating ALL correct answers.
        -   **Short-answer:** ~20% of questions. Requires a concise text response.
        -   **Coding challenge:** ~10% of questions. Provide a clear problem statement and expected output/logic. (If this is too complex for 1.5-flash to reliably generate, favor more short-answer).
    2.  Ensure questions cover both theoretical understanding and practical application of the skills.
    3.  Assign a unique `question_id` (e.g., "q1", "q2") to each question.
    4.  For each multiple/single choice question, you MUST provide the `correct_answer_keys` (a list of option values that are correct). This is CRITICAL for automated grading.
    
    **JSON Output Schema (List of Question Objects):**
{ASSESSMENT_QUESTIONS_SCHEMA.strip()}
    **Critical Rules:**
    - Your final output MUST be a JSON array containing exactly {num_questions} question objects.
    - DO NOT include any introductory or concluding text outside the JSON.
    - Ensure `correct_answer_keys` is always a LIST, even if only one answer.
    """

    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        questions = _safe_json_loads(response.text, fallback=None)
        
        if not questions or not isinstance(questions, list):
            print("\n--- ERROR: GEMINI FLASH FAILED TO GENERATE VALID ASSESSMENT QUESTIONS ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("----------------------------------------------------------------------\n")
            return None
        
        return {"questions": questions}

    except Exception as e:
        print(f"Error during AI question generation with Gemini Flash: {e}")
        return None


def evaluate_assessment_answers(
    user_id: str,
    submitted_answers: List[Dict[str, Any]],
    # original_questions: List[Dict[str, Any]] = None
) -> Optional[Dict[str, Any]]:
    """
    Evaluates user's assessment answers using Gemini Flash and provides structured results.
    """

    answers_summary = []
    for ans in submitted_answers:
        q_id = ans.get("question_id", "N/A")
        user_response = ans.get("answer")
        
        if isinstance(user_response, list):
            user_response_str = ", ".join(user_response)
        elif user_response is None:
            user_response_str = "No answer provided"
        else:
            user_response_str = str(user_response)
        
        answers_summary.append(f"Question ID: {q_id}\nUser Answer: ```{user_response_str}```\n---")

    answers_text = "\n".join(answers_summary)

    prompt = f"""
    You are an expert technical interviewer and AI grader.
    Your task is to evaluate a user's submitted answers for a skill assessment.
    Provide a comprehensive, structured evaluation based on the answers provided.

    **Instructions for Evaluation:**
    1.  **Calculate Overall Score:** Assign an overall percentage score (0-100%) for the assessment.
    2.  **Identify Skills Mastered/Areas to Improve (Counts):** Based on the questions and answers, estimate how many distinct skills were demonstrated proficiently and how many need significant improvement.
    3.  **List Strengths:** Provide 2-3 specific bullet points highlighting what the user did well.
    4.  **List Weaknesses:** Provide 2-3 specific bullet points highlighting areas where the user struggled or demonstrated gaps.
    5.  **Personalized Recommendations:** Provide 2-3 actionable, general recommendations for improvement. These should be text-based recommendations, not URLs.

    **User's Submitted Answers:**
    {'-'*30}
    {answers_text}
    {'-'*30}

    **JSON Output Schema:**
{ASSESSMENT_EVALUATION_SCHEMA.strip()}
    **Critical Rules:**
    - Your final output MUST be a single, valid JSON object following the schema.
    - DO NOT include any introductory or concluding text outside the JSON.
    - The `skill_scores` should be an object mapping skill names (e.g., Python, SQL) to a proficiency score (0-100). Infer these skills from the context of the assessment.
    """

    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        results = _safe_json_loads(response.text, fallback=None)
        
        if not results or not isinstance(results, dict):
            print("\n--- ERROR: GEMINI FLASH FAILED TO EVALUATE ASSESSMENT ANSWERS ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-------------------------------------------------------------------\n")
            return None
        
        return results

    except Exception as e:
        print(f"Error during AI assessment evaluation with Gemini Flash: {e}")
        return None
    

def generate_full_resume_analysis(resume_text: str, job_description: Optional[str] = None) -> Optional[Dict[str, Any]]:
    """
    Generates a comprehensive resume analysis report, including overall score,
    ATS score, strengths, areas for improvement, and section-wise feedback.
    """
    job_desc_context = ""
    job_role_hint = "General Candidate" # Default value
    if job_description and job_description.strip():
        job_desc_context = f"""
    The user has provided a job description. Analyze the resume specifically against this job description
    to provide highly tailored feedback, especially for ATS optimization, strengths, and areas for improvement.
    Infer the primary 'Job Role' from this description.

    **Job Description:**
    ```
    {job_description}
    ```
    """
        # A simple AI call to infer job role for display purposes
        model_name_for_role_inference = genai.GenerativeModel(MODEL_NAME)
        try:
            safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
            role_response = model_name_for_role_inference.generate_content(
                f"Extract the primary job role from the following job description. Respond with only the job role text (e.g., 'Software Engineer', 'Data Scientist', 'Frontend Developer').\n\nJob Description: {job_description}",
                safety_settings=safety_settings
            )
            inferred_role = role_response.text.strip()
            if inferred_role and len(inferred_role.split()) < 5: # Basic check to ensure it's a role, not a sentence
                job_role_hint = inferred_role
        except Exception as e:
            print(f"Warning: Could not infer job role from JD. Using default. Error: {e}")


    prompt = f"""
    You are an expert HR consultant and AI resume analyst. Your task is to provide a comprehensive analysis of the given resume.
    Generate a detailed report covering overall assessment, specific section analyses, key strengths, areas for improvement,
    and a dedicated ATS optimization score, all in a single JSON object.

    **Instructions:**
    1.  **Analysis Date:** Current date (e.g., "September 05, 2025").
    2.  **Job Role Context:** Infer a primary job role from the provided job description (if any) or from the resume itself. Default to "General Candidate" if unclear.
    3.  **AI Model:** "Google Gemini"
    4.  **Overall Resume Score:** A percentage (0-100) reflecting general quality, clarity, and effectiveness.
    5.  **Overall Resume Grade:** A concise word (e.g., "Excellent", "Good", "Fair", "Needs Improvement") corresponding to the score.
    6.  **ATS Optimization Score:** A percentage (0-100) reflecting compatibility with Applicant Tracking Systems, especially considering the job description.
    7.  **Section-wise Analysis:** Provide a 'title' and 'summary' for:
        -   `professional_profile_analysis`: For the summary/objective section.
        -   `education_analysis`: For the education section.
        -   `experience_analysis`: For work experience and projects.
        -   `skills_analysis`: For the skills section.
    8.  **Key Strengths:** 2-3 bullet points highlighting positive aspects.
    9.  **Areas for Improvement:** 3-5 bullet points covering general resume improvements AND specific ATS issues (e.g., keyword gaps, formatting problems).
    10. **Overall Assessment:** A concluding paragraph summarizing the findings and potential for improvement.

    {job_desc_context}

    **Resume Text:**
    ```
    {resume_text}
    ```

    **JSON Output Schema:**
{FULL_RESUME_ANALYSIS_SCHEMA.strip()}
    **Critical Rules:**
    - Your final output MUST be a single, valid JSON object following the schema.
    - DO NOT include any introductory or concluding text outside the JSON.
    - Ensure all scores are integers (0-100).
    - If no job description is provided, make reasonable general assumptions for the 'Job Role Context' and ATS analysis.
    - For `analysis_date`, always use the current date in 'Month DD, YYYY' format.
    - For section summaries, be direct and actionable, similar to the provided examples.
    """
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        analysis_data = _safe_json_loads(response.text, fallback=None)
        
        if not analysis_data or not isinstance(analysis_data, dict):
            print("\n--- ERROR: GEMINI FAILED TO GENERATE VALID FULL RESUME ANALYSIS ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("------------------------------------------------------------------\n")
            return None
        
        # Override job_role_context if successfully inferred
        if job_role_hint != "General Candidate": 
            analysis_data['job_role_context'] = job_role_hint
        
        # Ensure analysis_date is always current
        analysis_data['analysis_date'] = datetime.now().strftime("%B %d, %Y")

        return analysis_data
    except Exception as e:
        print(f"Error during Full Resume Analysis with Gemini: {e}")
        return None
    

# =======================================================================
# --- THIS IS THE FUNCTION FOR THE MOCK INTERVIEW CHAT ---
# =======================================================================
def get_interview_chat_response(job_description: str, history: List[Dict[str, str]], difficulty: str) -> Optional[Dict[str, str]]:
    """
    Acts as an AI Interviewer with adjustable difficulty.
    """
    # (The existing logic for this function remains unchanged)
    if difficulty == 'easy':
        personality_prompt = """
        Your Persona: You are a friendly and encouraging hiring manager for an entry-level role.
        Your Goal: Understand the candidate's basic knowledge and potential. Ask foundational, single-topic conceptual questions (e.g., "In Python, what is the difference between a list and a tuple?").
        Your Tone: Supportive and patient.
        Your First Action: Start with a simple, welcoming question like "Thanks for coming in. To start, could you tell me about a project you're proud of that's relevant to this role?"
        """
    elif difficulty == 'hard':
        personality_prompt = """
        Your Persona: You are a sharp, direct senior engineer conducting a final-round interview.
        Your Goal: Rigorously test the candidate's deep technical expertise, problem-solving, and system design skills. Ask challenging, multi-part, or scenario-based questions (e.g., "Given the requirements in the job description, walk me through how you would design a scalable, resilient API for our service. What bottlenecks would you anticipate and how would you mitigate them?").
        Your Tone: Critical, professional, and expecting detailed answers. You will ask tough follow-up questions.
        Your First Action: Start directly with a challenging technical question based on a core skill from the job description.
        """
    else: # Default to Medium
        personality_prompt = """
        Your Persona: You are a professional team lead for a mid-level role.
        Your Goal: Assess the candidate's practical skills and real-world project experience. Ask behavioral and technical questions that require specific examples (e.g., "Tell me about a time you had to deal with significant technical debt. How did you handle it and what was the outcome?").
        Your Tone: Objective and focused.
        Your First Action: Start with a question about the candidate's most relevant experience from their resume, tying it to the job description.
        """

    formatted_history = [{'role': msg['role'], 'parts': [{'text': msg['content']}]} for msg in history]
    system_instruction = f"""
    {personality_prompt}
    
    CRITICAL RULE: You are the INTERVIEWER. The user is the CANDIDATE. You must conduct a realistic interview.
    Base ALL of your questions and analysis strictly on the provided job description context below. Do not ask about skills not mentioned.

    --- JOB DESCRIPTION CONTEXT ---
    {job_description}
    --- END CONTEXT ---
    """
    full_history = [
        {'role': 'user', 'parts': [{'text': system_instruction}]},
        {'role': 'model', 'parts': [{'text': "Understood. I am ready to begin the interview."}]}
    ] + formatted_history
    
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        chat_session = model.start_chat(history=full_history[:-1])
        last_user_message = full_history[-1]['parts'][0]['text']
        response = chat_session.send_message(
            last_user_message,
            generation_config=genai.types.GenerationConfig(temperature=0.7),
        )
        return {"reply": response.text}
    except Exception as e:
        print(f"An error occurred in the interview chat endpoint: {e}")
        return None

# =======================================================================
# --- NEW FUNCTION FOR INTERVIEW SUMMARY ---
# =======================================================================
def get_interview_summary(job_description: str, history: List[Dict[str, str]]) -> Optional[Dict[str, Any]]:
    """
    Analyzes the full interview transcript and provides a performance summary.
    """
    transcript = "\n".join([f"{msg['role']}: {msg['content']}" for msg in history])

    prompt = f"""
    You are an expert career coach and technical recruiter. Your task is to analyze the following mock interview transcript and provide a performance summary.
    
    **Job Description Context:**
    ```
    {job_description}
    ```

    **Interview Transcript:**
    ```
    {transcript}
    ```

    **Your Analysis Task:**
    Based on the job description and the transcript, provide a detailed analysis in a valid JSON object. The JSON must have the following keys:
    1.  `"overall_score"`: An integer from 0 to 100 representing the candidate's overall performance.
    2.  `"strengths"`: A list of 2-3 specific, positive points about the candidate's performance, citing examples from the transcript.
    3.  `"areas_for_improvement"`: A list of 2-3 specific, constructive points for improvement, citing examples.
    4.  `"overall_feedback"`: A concise paragraph summarizing the performance and providing a final recommendation.

    **Critical Rules:**
    - Your final output must be ONLY the valid JSON object. Do not include markdown or any other text.
    - Be honest and constructive in your feedback.
    """
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        summary_data = _safe_json_loads(response.text, fallback=None)
        
        if not summary_data:
            print("\n--- ERROR: GEMINI FAILED TO GENERATE VALID INTERVIEW SUMMARY ---")
            print("API Response Text:", response.text)
            return None

        return summary_data
    except Exception as e:
        print(f"Error generating interview summary: {e}")
        return None
